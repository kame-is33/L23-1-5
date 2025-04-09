"""
このファイルは、最初の画面読み込み時にのみ実行される初期化処理が記述されたファイルです。
"""

############################################################
# ライブラリの読み込み
############################################################
import os
import logging
from logging.handlers import TimedRotatingFileHandler
from uuid import uuid4
import sys
import unicodedata
from dotenv import load_dotenv
import streamlit as st
from docx import Document
from langchain_community.document_loaders import WebBaseLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
import constants as ct


############################################################
# 設定関連
############################################################
# 「.env」ファイルで定義した環境変数の読み込み
load_dotenv()


############################################################
# 関数定義
############################################################

def initialize():
    """
    画面読み込み時に実行する初期化処理
    """
    # 初期化データの用意
    initialize_session_state()
    # ログ出力用にセッションIDを生成
    initialize_session_id()
    # ログ出力の設定
    initialize_logger()
    # RAGのRetrieverを作成
    initialize_retriever()


def initialize_logger():
    """
    ログ出力の設定
    """
    try:
        # 指定のログフォルダが存在すれば読み込み、存在しなければ新規作成
        os.makedirs(ct.LOG_DIR_PATH, exist_ok=True)
        
        # 引数に指定した名前のロガー（ログを記録するオブジェクト）を取得
        # 再度別の箇所で呼び出した場合、すでに同じ名前のロガーが存在していれば読み込む
        logger = logging.getLogger(ct.LOGGER_NAME)

        # すでにロガーにハンドラー（ログの出力先を制御するもの）が設定されている場合、同じログ出力が複数回行われないよう処理を中断する
        if logger.hasHandlers():
            return

        # 1日単位でログファイルの中身をリセットし、切り替える設定
        log_handler = TimedRotatingFileHandler(
            os.path.join(ct.LOG_DIR_PATH, ct.LOG_FILE),
            when="D",
            encoding="utf8"
        )
        # 出力するログメッセージのフォーマット定義
        # - 「levelname」: ログの重要度（INFO, WARNING, ERRORなど）
        # - 「asctime」: ログのタイムスタンプ（いつ記録されたか）
        # - 「lineno」: ログが出力されたファイルの行番号
        # - 「funcName」: ログが出力された関数名
        # - 「session_id」: セッションID（誰のアプリ操作か分かるように）
        # - 「message」: ログメッセージ
        formatter = logging.Formatter(
            f"[%(levelname)s] %(asctime)s line %(lineno)s, in %(funcName)s, session_id={st.session_state.session_id}: %(message)s"
        )

        # 定義したフォーマッターの適用
        log_handler.setFormatter(formatter)

        # ログレベルを「INFO」に設定
        logger.setLevel(logging.INFO)

        # 作成したハンドラー（ログ出力先を制御するオブジェクト）を、
        # ロガー（ログメッセージを実際に生成するオブジェクト）に追加してログ出力の最終設定
        logger.addHandler(log_handler)
    except Exception as e:
        st.error(f"ログ設定エラー: {e}")


def initialize_session_id():
    """
    セッションIDの作成
    """
    if "session_id" not in st.session_state:
        # ランダムな文字列（セッションID）を、ログ出力用に作成
        st.session_state.session_id = uuid4().hex


def initialize_retriever():
    """
    画面読み込み時にRAGのRetriever（ベクターストアから検索するオブジェクト）を作成
    """
    # ロガーを読み込むことで、後続の処理中に発生したエラーなどがログファイルに記録される
    logger = logging.getLogger(ct.LOGGER_NAME)

    # すでにRetrieverが作成済みの場合、後続の処理を中断
    if "retriever" in st.session_state:
        return
    
    try:
        # RAGの参照先となるデータソースの読み込み
        logger.info("データソースの読み込みを開始")
        docs_all = load_data_sources()
        logger.info(f"{len(docs_all)}件のドキュメントを読み込みました")

        # OSがWindowsの場合、Unicode正規化と、cp932（Windows用の文字コード）で表現できない文字を除去
        for doc in docs_all:
            doc.page_content = adjust_string(doc.page_content)
            for key in doc.metadata:
                doc.metadata[key] = adjust_string(doc.metadata[key])
        
        # 埋め込みモデルの用意
        logger.info("埋め込みモデルの初期化")
        embeddings = OpenAIEmbeddings()
        
        # チャンク分割用のオブジェクトを作成
        text_splitter = CharacterTextSplitter(
            chunk_size=ct.CHUNK_SIZE,
            chunk_overlap=ct.CHUNK_OVERLAP,
            separator="\n"
        )
        
        # チャンク分割を実施
        logger.info("チャンク分割を実行")
        splitted_docs = text_splitter.split_documents(docs_all)
        logger.info(f"{len(splitted_docs)}件のチャンクに分割されました")
        
        # ベクターストアの作成
        logger.info("ベクターストアの作成")
        db = Chroma.from_documents(splitted_docs, embedding=embeddings)
        
        # ベクターストアを検索するRetrieverの作成
        logger.info(f"Retrieverの作成 (k={ct.RETRIEVER_DOCUMENT_COUNT})")
        st.session_state.retriever = db.as_retriever(search_kwargs={"k": ct.RETRIEVER_DOCUMENT_COUNT})
        logger.info("Retrieverの初期化完了")
    except Exception as e:
        logger.error(f"Retriever初期化エラー: {e}")
        raise


def initialize_session_state():
    """
    初期化データの用意
    """
    if "messages" not in st.session_state:
        # 「表示用」の会話ログを順次格納するリストを用意
        st.session_state.messages = []
        # 「LLMとのやりとり用」の会話ログを順次格納するリストを用意
        st.session_state.chat_history = []
    
    # 開発者モードのフラグ
    if "debug_mode" not in st.session_state:
        st.session_state.debug_mode = False


def load_data_sources():
    """
    RAGの参照先となるデータソースの読み込み

    Returns:
        読み込んだ通常データソース
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    
    # データソースを格納する用のリスト
    docs_all = []
    
    try:
        # ファイル読み込みの実行（渡した各リストにデータが格納される）
        recursive_file_check(ct.RAG_TOP_FOLDER_PATH, docs_all)
        logger.info(f"ファイルから{len(docs_all)}件のドキュメントを読み込みました")
        
        web_docs_all = []
        # ファイルとは別に、指定のWebページ内のデータも読み込み
        # 読み込み対象のWebページ一覧に対して処理
        for web_url in ct.WEB_URL_LOAD_TARGETS:
            try:
                # 指定のWebページを読み込み
                logger.info(f"Webページの読み込み: {web_url}")
                loader = WebBaseLoader(web_url)
                web_docs = loader.load()
                # for文の外のリストに読み込んだデータソースを追加
                web_docs_all.extend(web_docs)
            except Exception as e:
                logger.warning(f"Webページの読み込みエラー: {web_url} - {e}")
                
        # 通常読み込みのデータソースにWebページのデータを追加
        logger.info(f"Webから{len(web_docs_all)}件のドキュメントを読み込みました")
        docs_all.extend(web_docs_all)
    
    except Exception as e:
        logger.error(f"データソース読み込みエラー: {e}")
        raise

    return docs_all


def recursive_file_check(path, docs_all):
    """
    RAGの参照先となるデータソースの読み込み

    Args:
        path: 読み込み対象のファイル/フォルダのパス
        docs_all: データソースを格納する用のリスト
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    
    try:
        # パスがフォルダかどうかを確認
        if os.path.isdir(path):
            # フォルダの場合、フォルダ内のファイル/フォルダ名の一覧を取得
            files = os.listdir(path)
            # 各ファイル/フォルダに対して処理
            for file in files:
                # ファイル/フォルダ名だけでなく、フルパスを取得
                full_path = os.path.join(path, file)
                # フルパスを渡し、再帰的にファイル読み込みの関数を実行
                recursive_file_check(full_path, docs_all)
        else:
            # パスがファイルの場合、ファイル読み込み
            file_load(path, docs_all)
    except PermissionError:
        logger.warning(f"アクセス権限がありません: {path}")
    except Exception as e:
        logger.error(f"ファイル読み込みエラー: {path} - {e}")


def file_load(path, docs_all):
    """
    ファイル内のデータ読み込み

    Args:
        path: ファイルパス
        docs_all: データソースを格納する用のリスト
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    
    try:
        # ファイルの拡張子を取得
        file_extension = os.path.splitext(path)[1].lower()
        # ファイル名（拡張子を含む）を取得
        file_name = os.path.basename(path)

        # 想定していたファイル形式の場合のみ読み込む
        if file_extension in ct.SUPPORTED_EXTENSIONS:
            logger.info(f"ファイル読み込み: {path}")
            # ファイルの拡張子に合ったdata loaderを使ってデータ読み込み
            loader = ct.SUPPORTED_EXTENSIONS[file_extension](path)
            docs = loader.load()
            docs_all.extend(docs)
        elif file_extension == ".docx":
            # .docxファイルの場合、専用の処理
            text = extract_docx_text(path)
            from langchain.schema import Document as LangchainDoc
            doc = LangchainDoc(page_content=text, metadata={"source": path})
            docs_all.append(doc)
            logger.info(f"DOCXファイル読み込み: {path}")
    except Exception as e:
        logger.warning(f"ファイル読み込みスキップ: {path} - {e}")


def adjust_string(s):
    """
    Windows環境でRAGが正常動作するよう調整
    
    Args:
        s: 調整を行う文字列
    
    Returns:
        調整を行った文字列
    """
    # 調整対象は文字列のみ
    if type(s) is not str:
        return s

    # OSがWindowsの場合、Unicode正規化と、cp932（Windows用の文字コード）で表現できない文字を除去
    if sys.platform.startswith("win"):
        s = unicodedata.normalize('NFC', s)
        s = s.encode("cp932", "ignore").decode("cp932")
        return s
    
    # OSがWindows以外の場合はそのまま返す
    return s


def extract_docx_text(path):
    """
    Word文書（.docx）から段落・見出しを抽出して結合したテキストを返す

    Args:
        path: 読み込むdocxファイルのパス

    Returns:
        str: 結合されたテキスト
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    
    try:
        doc = Document(path)
        paragraphs = []
        
        # 段落の抽出
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # 表の内容抽出
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_texts.append(cell.text.strip())
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))
        
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX抽出エラー: {path} - {e}")
        return ""